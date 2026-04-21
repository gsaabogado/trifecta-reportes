"""
Trifecta CCI Report Generator
Generates a Customs Clearance Inspection report (.docx) from a folder of photos.

Usage:
    python generate_report.py /path/to/inspection/folder [-o output.docx]
"""

import argparse
import os
import re
import sys
from pathlib import Path
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from PIL import Image

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
SCRIPT_DIR = Path(__file__).resolve().parent
PLANTILLAS_DIR = SCRIPT_DIR.parent / "plantillas"
LOGO_PATH = PLANTILLAS_DIR / "logo_trifecta.png"

# Trifecta brand colors (from Krammer report)
NAVY = RGBColor(0x00, 0x2E, 0x6D)       # #002E6D — headers, accents
CHARCOAL = RGBColor(0x2D, 0x2A, 0x26)    # #2D2A26 — body text
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)
MID_GRAY = RGBColor(0xBF, 0xBF, 0xBF)
PLACEHOLDER_GRAY = RGBColor(0xA0, 0xA0, 0xA0)

FONT_NAME = "Calibri"
BODY_SIZE = Pt(10)
SMALL_SIZE = Pt(8)
CAPTION_SIZE = Pt(8)

# Page layout
PAGE_WIDTH_IN = 8.5
PAGE_HEIGHT_IN = 11.0

# ---------------------------------------------------------------------------
# Translations
# ---------------------------------------------------------------------------
STRINGS = {
    "en": {
        "title": "Customs Clearance Inspection Report",
        "subtitle": "Pre-Shipment Container Loading Inspection",
        "client": "Client:",
        "product": "Product:",
        "container": "Container:",
        "inspection_date": "Inspection date:",
        "inspector": "Inspector:",
        "cci_reference": "CCI reference:",
        "sec_summary": "1. Inspection Summary",
        "chk_header_item": "Inspection Item",
        "chk_header_result": "Result",
        "chk_header_comments": "Comments",
        "chk_qty_packing": "Quantity matches Packing List",
        "chk_num_items": "Number of items per requirement",
        "chk_packing_cond": "Packing conditions and shipping marks",
        "chk_shipping_marks": "Shipping marks on boxes",
        "chk_labeling": "Product labeling",
        "chk_container_cond": "Container conditions",
        "chk_loading": "Loading process matches packing description",
        "chk_fumigation": "Fumigation certificate",
        "chk_nom": "NOM compliance",
        "sec_order_details": "2. Order Details",
        "po_number": "PO number:",
        "order_quantity": "Order quantity:",
        "total_cartons": "Total cartons:",
        "supplier": "Supplier:",
        "contact_person": "Contact person:",
        "weather": "Weather:",
        "start_time": "Start time:",
        "remarks": "Remarks",
        "no_remarks": "No remarks.",
        "sec_packing_list": "3. Product Packing List",
        "pl_item_no": "Item No.",
        "pl_description": "Description",
        "pl_qty_ctns": "Qty (Ctns)",
        "pl_total_pcs": "Total Pcs",
        "pl_loading_ctns": "Loading Ctns",
        "pl_final_qty": "Final Loading Quantity",
        "sec_product_details": "4. Product Details and Pictures",
        "label": "Label:",
        "result": "Result:",
        "general_pictures": "General Product Pictures",
        "product_caption": "product",
        "sec_container": "5. Container Details",
        "ct_type": "Container Type",
        "ct_number": "Number",
        "ct_seal": "Seal No.",
        "ct_loading_qty": "Loading Quantity",
        "cap_container_num": "Container Number",
        "cap_seal": "Seal No.",
        "cap_empty": "Empty Container",
        "cap_light_test": "Light Test (with Flash)",
        "cap_container": "Container",
        "container_damages": "Container Damages",
        "placeholder_damages": "Insert container damage photos here if applicable.",
        "sec_loading": "6. Loading Process",
        "placeholder_photo": "[Insert photo]",
        "sec_acknowledgment": "7. CCI Acknowledgment",
        "placeholder_ack": "Insert scanned CCI acknowledgment document here.",
        "ack_title": "CUSTOMS CLEARANCE INSPECTION ACKNOWLEDGMENT",
        "ack_title_zh": "工厂告知书",
        "ack_date": "Date 日期:",
        "ack_container_no": "Container No. 集装箱号:",
        "ack_product": "Product 产品:",
        "ack_items": "Items 款式数量:",
        "ack_loading_time": "Loading time 装货起始时间:",
        "ack_reference_no": "Reference No. 编号:",
        "ack_seal_no": "Seal No. 铅封号:",
        "ack_sample_standard": "Sample standard 抽样标准:",
        "ack_package_qty": "Package quantity 包装数量:",
        "ack_photos_allowed": "Supplier allowed to take pictures 是否允许拍照",
        "ack_labels": "Label pictures 标签照片",
        "ack_fumigation": "Fumigation stamp 熏蒸标志",
        "ack_qty_check": "Quantity check 数量检查",
        "ack_container_quality": "Container quality 集装箱质量",
        "ack_box_quality": "Box quality 包装质量",
        "ack_all_items_photo": "All items were photographed 所有产品照片记录",
        "ack_space_enough": "Space was enough for merchandise 空间足够装所有货",
        "ack_same_qty_pl": "Same quantity as Packing List 数量与箱单一致",
        "ack_inspection_result": "Inspection result 检验结果:",
        "ack_pass": "PASS 合格",
        "ack_fail": "FAIL 不合格",
        "ack_disclaimer": "This report documents the conditions of the shipment at the time of inspection. Our inspectors are not authorized to receive money or gifts that may influence the inspection result. Our inspectors cannot make decisions about the shipment (restrictions and authorizations). In case of any anomaly, the inspector must inform Trifecta offices directly.\n本报告仅告知装货时的实际情况。检查时间内，我们的检验员不允许接收任何可能影响检验结果的费用和礼物。我们的检验员不允许做出关于货物（限制和许可）的任何决定。如有任何异常，检验员必须直接通知Trifecta办事处。",
        "ack_comments": "Comments and remarks 备注:",
        "ack_inspector": "INSPECTOR 检查员",
        "ack_supplier": "SUPPLIER 供应商",
    },
    "es": {
        "title": "Reporte de Previo en Origen",
        "subtitle": "Inspección de Carga de Contenedor Pre-Embarque",
        "client": "Cliente:",
        "product": "Producto:",
        "container": "Contenedor:",
        "inspection_date": "Fecha de inspección:",
        "inspector": "Inspector:",
        "cci_reference": "Referencia CCI:",
        "sec_summary": "1. Resumen de Inspección",
        "chk_header_item": "Concepto de Inspección",
        "chk_header_result": "Resultado",
        "chk_header_comments": "Comentarios",
        "chk_qty_packing": "Cantidad coincide con Lista de Empaque",
        "chk_num_items": "Número de artículos según requerimiento",
        "chk_packing_cond": "Condiciones de empaque y marcas de envío",
        "chk_shipping_marks": "Marcas de envío en cajas",
        "chk_labeling": "Etiquetado del producto",
        "chk_container_cond": "Condiciones del contenedor",
        "chk_loading": "Proceso de carga conforme a descripción de empaque",
        "chk_fumigation": "Certificado de fumigación",
        "chk_nom": "Cumplimiento de NOM",
        "sec_order_details": "2. Detalles del Pedido",
        "po_number": "Número de PO:",
        "order_quantity": "Cantidad del pedido:",
        "total_cartons": "Total de cajas:",
        "supplier": "Proveedor:",
        "contact_person": "Persona de contacto:",
        "weather": "Clima:",
        "start_time": "Hora de inicio:",
        "remarks": "Observaciones",
        "no_remarks": "Sin observaciones.",
        "sec_packing_list": "3. Lista de Empaque",
        "pl_item_no": "No. de Artículo",
        "pl_description": "Descripción",
        "pl_qty_ctns": "Cant. (Cajas)",
        "pl_total_pcs": "Total Pzas",
        "pl_loading_ctns": "Cajas Cargadas",
        "pl_final_qty": "Cantidad Final de Carga",
        "sec_product_details": "4. Detalles y Fotografías del Producto",
        "label": "Etiqueta:",
        "result": "Resultado:",
        "general_pictures": "Fotografías Generales del Producto",
        "product_caption": "producto",
        "sec_container": "5. Detalles del Contenedor",
        "ct_type": "Tipo de Contenedor",
        "ct_number": "Número",
        "ct_seal": "No. de Sello",
        "ct_loading_qty": "Cantidad de Carga",
        "cap_container_num": "Número de Contenedor",
        "cap_seal": "No. de Sello",
        "cap_empty": "Contenedor Vacío",
        "cap_light_test": "Prueba de Luz (con Flash)",
        "cap_container": "Contenedor",
        "container_damages": "Daños al Contenedor",
        "placeholder_damages": "Insertar fotografías de daños al contenedor si aplica.",
        "sec_loading": "6. Proceso de Carga",
        "placeholder_photo": "[Insertar foto]",
        "sec_acknowledgment": "7. Acuse de Recibo CCI",
        "placeholder_ack": "Insertar documento de acuse de recibo CCI escaneado aquí.",
        "ack_title": "ACUSE DE RECIBO DE INSPECCIÓN DE DESPACHO ADUANAL",
        "ack_title_zh": "工厂告知书",
        "ack_date": "Fecha 日期:",
        "ack_container_no": "No. de Contenedor 集装箱号:",
        "ack_product": "Producto 产品:",
        "ack_items": "Artículos 款式数量:",
        "ack_loading_time": "Hora de carga 装货起始时间:",
        "ack_reference_no": "No. de Referencia 编号:",
        "ack_seal_no": "No. de Sello 铅封号:",
        "ack_sample_standard": "Estándar de muestreo 抽样标准:",
        "ack_package_qty": "Cantidad de paquetes 包装数量:",
        "ack_photos_allowed": "Proveedor permitió tomar fotografías 是否允许拍照",
        "ack_labels": "Fotografías de etiquetas 标签照片",
        "ack_fumigation": "Sello de fumigación 熏蒸标志",
        "ack_qty_check": "Verificación de cantidad 数量检查",
        "ack_container_quality": "Calidad del contenedor 集装箱质量",
        "ack_box_quality": "Calidad de las cajas 包装质量",
        "ack_all_items_photo": "Todos los artículos fueron fotografiados 所有产品照片记录",
        "ack_space_enough": "Espacio suficiente para la mercancía 空间足够装所有货",
        "ack_same_qty_pl": "Misma cantidad que la Lista de Empaque 数量与箱单一致",
        "ack_inspection_result": "Resultado de la inspección 检验结果:",
        "ack_pass": "APROBADO 合格",
        "ack_fail": "RECHAZADO 不合格",
        "ack_disclaimer": "Este reporte documenta las condiciones del embarque al momento de la inspección. Nuestros inspectores no están autorizados a recibir dinero ni obsequios que puedan influir en el resultado de la inspección. Nuestros inspectores no pueden tomar decisiones sobre el embarque (restricciones y autorizaciones). En caso de cualquier anomalía, el inspector debe informar directamente a las oficinas de Trifecta.\n本报告仅告知装货时的实际情况。检查时间内，我们的检验员不允许接收任何可能影响检验结果的费用和礼物。我们的检验员不允许做出关于货物（限制和许可）的任何决定。如有任何异常，检验员必须直接通知Trifecta办事处。",
        "ack_comments": "Comentarios y observaciones 备注:",
        "ack_inspector": "INSPECTOR 检查员",
        "ack_supplier": "PROVEEDOR 供应商",
    },
}


def t(key, lang="en"):
    """Get translated string."""
    return STRINGS.get(lang, STRINGS["en"]).get(key, STRINGS["en"].get(key, key))


# ---------------------------------------------------------------------------
# Helpers
# Form field counter for unique names
_field_counter = [0]


def add_form_field(paragraph, field_name=None, default_text="", size=BODY_SIZE, color=CHARCOAL):
    """Insert a Word legacy text form field into a paragraph.
    These become fillable fields when exported to PDF."""
    _field_counter[0] += 1
    name = field_name or f"field_{_field_counter[0]}"

    # Begin field char with form field data
    fld_begin = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'  <w:rPr><w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/>'
        f'    <w:sz w:val="{int(size.pt * 2)}"/>'
        f'    <w:color w:val="{str(color)}"/>'
        f'  </w:rPr>'
        f'  <w:fldChar w:fldCharType="begin">'
        f'    <w:ffData>'
        f'      <w:name w:val="{name}"/>'
        f'      <w:enabled/>'
        f'      <w:calcOnExit w:val="0"/>'
        f'      <w:textInput>'
        f'        <w:default w:val="{default_text}"/>'
        f'      </w:textInput>'
        f'    </w:ffData>'
        f'  </w:fldChar>'
        f'</w:r>'
    )
    # Instruction text
    fld_instr = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'  <w:rPr><w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/>'
        f'    <w:sz w:val="{int(size.pt * 2)}"/>'
        f'    <w:color w:val="{str(color)}"/>'
        f'  </w:rPr>'
        f'  <w:instrText xml:space="preserve"> FORMTEXT </w:instrText>'
        f'</w:r>'
    )
    # Separator
    fld_sep = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'  <w:rPr><w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/>'
        f'    <w:sz w:val="{int(size.pt * 2)}"/>'
        f'    <w:color w:val="{str(color)}"/>'
        f'  </w:rPr>'
        f'  <w:fldChar w:fldCharType="separate"/>'
        f'</w:r>'
    )
    # Default text display
    display = default_text if default_text else "\u200B"  # zero-width space if empty
    fld_text = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'  <w:rPr><w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/>'
        f'    <w:sz w:val="{int(size.pt * 2)}"/>'
        f'    <w:color w:val="{str(color)}"/>'
        f'    <w:noProof/>'
        f'  </w:rPr>'
        f'  <w:t xml:space="preserve">{display}</w:t>'
        f'</w:r>'
    )
    # End field char
    fld_end = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'  <w:rPr><w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/>'
        f'    <w:sz w:val="{int(size.pt * 2)}"/>'
        f'    <w:color w:val="{str(color)}"/>'
        f'  </w:rPr>'
        f'  <w:fldChar w:fldCharType="end"/>'
        f'</w:r>'
    )

    p_elem = paragraph._p
    p_elem.append(fld_begin)
    p_elem.append(fld_instr)
    p_elem.append(fld_sep)
    p_elem.append(fld_text)
    p_elem.append(fld_end)


def add_form_field_to_cell(cell, field_name=None, default_text="",
                           size=Pt(10), color=CHARCOAL,
                           align=WD_ALIGN_PARAGRAPH.LEFT):
    """Add a fillable form field inside a table cell."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    add_form_field(p, field_name, default_text, size, color)


# ---------------------------------------------------------------------------
def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set cell borders. kwargs: top, bottom, left, right."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = parse_xml(f"<w:tcBorders {nsdecls('w')}/>")
        tcPr.append(tcBorders)
    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{attrs.get("val", "single")}" '
            f'w:sz="{attrs.get("sz", "4")}" w:space="0" '
            f'w:color="{attrs.get("color", "002E6D")}"/>'
        )
        existing = tcBorders.find(qn(f"w:{edge}"))
        if existing is not None:
            tcBorders.remove(existing)
        tcBorders.append(element)


def set_cell_no_borders(cell):
    """Remove all borders from a cell."""
    for edge in ("top", "bottom", "left", "right"):
        set_cell_border(cell, **{edge: {"val": "none", "sz": "0", "color": "FFFFFF"}})


def set_cell_margins(cell, top=0, bottom=0, left=80, right=80):
    """Set cell margins in twips."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    existing = tcPr.find(qn("w:tcMar"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcMar)


def fmt(cell, text, size=BODY_SIZE, bold=False, color=CHARCOAL,
        align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    """Write formatted text into a cell."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.italic = italic
    return run


def add_section_heading(doc, text):
    """Add an elegant section heading: navy text with thin rule below."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = NAVY
    # Thin bottom border on paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="3" w:color="002E6D"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def add_subsection_heading(doc, text):
    """Add a smaller subsection heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = NAVY
    return p


def add_body_text(doc, text, italic=False, color=CHARCOAL, size=BODY_SIZE):
    """Add a body paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.color.rgb = color
    run.font.italic = italic
    return p


def add_placeholder(doc, text):
    """Add a gray placeholder text."""
    return add_body_text(doc, text, italic=True, color=PLACEHOLDER_GRAY, size=SMALL_SIZE)


def resize_image_for_doc(image_path, max_width_inches=3.3, max_height_inches=2.5, quality=85):
    """Resize and compress an image, return BytesIO buffer."""
    img = Image.open(image_path)
    if img.mode in ("RGBA", "P"):
        img = img.convert("RGB")
    dpi = 300
    max_w_px = int(max_width_inches * dpi)
    max_h_px = int(max_height_inches * dpi)
    img.thumbnail((max_w_px, max_h_px), Image.LANCZOS)
    buf = BytesIO()
    img.save(buf, format="JPEG", quality=quality, optimize=True)
    buf.seek(0)
    return buf, img.size


def natural_sort_key(s):
    return [int(c) if c.isdigit() else c.lower() for c in re.split(r'(\d+)', str(s))]


def get_sorted_images(folder_path):
    extensions = {".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif"}
    images = [
        f for f in Path(folder_path).iterdir()
        if f.is_file() and f.suffix.lower() in extensions
    ]
    return sorted(images, key=lambda f: natural_sort_key(f.name))


def make_table_borderless(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


# ---------------------------------------------------------------------------
# Document construction
# ---------------------------------------------------------------------------
def create_document():
    """Create a new document with Trifecta branding and margins."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.page_width = Inches(PAGE_WIDTH_IN)
        section.page_height = Inches(PAGE_HEIGHT_IN)
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = BODY_SIZE
    font.color.rgb = CHARCOAL
    # Paragraph spacing
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)
    return doc


def add_header_footer(doc):
    """Add Trifecta logo in header (right-aligned) and contact info in footer."""
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    # Header — small logo, right-aligned
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if LOGO_PATH.exists():
        run = hp.add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(1.2))

    # Footer — thin rule + contact line
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add top border to footer paragraph
    pPr = fp._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="4" w:color="BFBFBF"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    run = fp.add_run("trifecta-inspections.com  |  Shenzhen  |  Guanajuato  |  Zurich  |  Prague")
    run.font.name = FONT_NAME
    run.font.size = Pt(7)
    run.font.color.rgb = MID_GRAY


def build_page1_summary(doc, lang="en"):
    """Page 1: Inspection summary — clean, professional layout."""
    # Title
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(t("title", lang))
    run.font.name = FONT_NAME
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = NAVY

    # Subtitle
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(t("subtitle", lang))
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = PLACEHOLDER_GRAY

    # Metadata block
    meta_fields = [
        (t("client", lang), ""),
        (t("product", lang), ""),
        (t("container", lang), ""),
        (t("inspection_date", lang), ""),
        (t("inspector", lang), ""),
        (t("cci_reference", lang), ""),
    ]
    table = doc.add_table(rows=len(meta_fields), cols=2)
    make_table_borderless(table)
    table.autofit = True
    for i, (label, default) in enumerate(meta_fields):
        row = table.rows[i]
        # Label
        cell_l = row.cells[0]
        cell_l.width = Inches(1.5)
        p = cell_l.paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(label)
        run.font.name = FONT_NAME
        run.font.size = Pt(9)
        run.font.color.rgb = PLACEHOLDER_GRAY
        # Value — fillable form field
        cell_v = row.cells[1]
        cell_v.width = Inches(4.0)
        field_key = label.replace(":", "").strip().replace(" ", "_").lower()
        add_form_field_to_cell(cell_v, field_name=f"meta_{field_key}",
                               size=Pt(10), color=CHARCOAL)

    # --- Inspection Summary ---
    add_section_heading(doc, t("sec_summary", lang))

    # Summary checklist table
    checklist_headers = [
        t("chk_header_item", lang),
        t("chk_header_result", lang),
        t("chk_header_comments", lang),
    ]
    checklist_items = [
        t("chk_qty_packing", lang),
        t("chk_num_items", lang),
        t("chk_packing_cond", lang),
        t("chk_shipping_marks", lang),
        t("chk_labeling", lang),
        t("chk_container_cond", lang),
        t("chk_loading", lang),
        t("chk_fumigation", lang),
        t("chk_nom", lang),
    ]

    table = doc.add_table(rows=1 + len(checklist_items), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row — navy background
    for j, header in enumerate(checklist_headers):
        cell = table.rows[0].cells[j]
        set_cell_shading(cell, "002E6D")
        fmt(cell, header, size=Pt(9), bold=True, color=WHITE,
            align=WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_border(cell, bottom={"sz": "6", "color": "002E6D"})

    # Data rows
    for i, item in enumerate(checklist_items):
        row = table.rows[i + 1]
        # Item name
        fmt(row.cells[0], item, size=Pt(9), color=CHARCOAL)
        # Result — fillable
        set_cell_shading(row.cells[1], "F2F2F2")
        add_form_field_to_cell(row.cells[1], field_name=f"chk_result_{i}",
                               size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER)
        # Comments — fillable
        add_form_field_to_cell(row.cells[2], field_name=f"chk_comment_{i}",
                               size=Pt(9))
        # Thin bottom border
        for j in range(3):
            set_cell_border(row.cells[j], bottom={"sz": "2", "color": "D9D9D9"})

    # Column widths
    widths = [Inches(3.5), Inches(1.0), Inches(2.0)]
    for row in table.rows:
        for j, w in enumerate(widths):
            row.cells[j].width = w

    # Remarks — fillable
    add_subsection_heading(doc, t("remarks", lang))
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    add_form_field(p, field_name="remarks_summary", size=Pt(9), color=PLACEHOLDER_GRAY)


def build_page2_packing_list(doc, lang="en"):
    """Page 2: Order details + packing list."""
    doc.add_page_break()

    # --- Order Details (moved here to avoid blank page) ---
    add_section_heading(doc, t("sec_order_details", lang))

    details_fields = [
        (t("po_number", lang), ""),
        (t("order_quantity", lang), ""),
        (t("total_cartons", lang), ""),
        (t("supplier", lang), ""),
        (t("contact_person", lang), ""),
        (t("weather", lang), ""),
        (t("start_time", lang), ""),
    ]
    table2 = doc.add_table(rows=len(details_fields), cols=2)
    make_table_borderless(table2)
    for i, (label, default) in enumerate(details_fields):
        row = table2.rows[i]
        cell_l = row.cells[0]
        cell_l.width = Inches(1.8)
        p = cell_l.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(label)
        run.font.name = FONT_NAME
        run.font.size = Pt(9)
        run.font.color.rgb = CHARCOAL
        run.font.bold = True

        cell_v = row.cells[1]
        cell_v.width = Inches(4.0)
        set_cell_border(cell_v, bottom={"sz": "2", "color": "D9D9D9"})
        field_key = label.replace(":", "").strip().replace(" ", "_").lower()
        add_form_field_to_cell(cell_v, field_name=f"detail_{field_key}",
                               size=Pt(10), color=CHARCOAL)

    # --- Packing List ---
    add_section_heading(doc, t("sec_packing_list", lang))

    headers = [
        t("pl_item_no", lang),
        t("pl_description", lang),
        t("pl_qty_ctns", lang),
        t("pl_total_pcs", lang),
        t("pl_loading_ctns", lang),
    ]
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        set_cell_shading(cell, "002E6D")
        fmt(cell, header, size=Pt(9), bold=True, color=WHITE,
            align=WD_ALIGN_PARAGRAPH.CENTER)

    # 5 empty rows
    for _ in range(5):
        row = table.add_row()
        for j in range(5):
            cell = row.cells[j]
            align = WD_ALIGN_PARAGRAPH.CENTER if j >= 2 else WD_ALIGN_PARAGRAPH.LEFT
            add_form_field_to_cell(cell, field_name=f"pl_r{_}_{j}",
                                   size=Pt(9), align=align)
            set_cell_border(cell, bottom={"sz": "2", "color": "D9D9D9"})

    # Total row
    row = table.add_row()
    row.cells[0].merge(row.cells[3])
    cell = row.cells[0]
    fmt(cell, t("pl_final_qty", lang), size=Pt(9), bold=True, color=WHITE,
        align=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell_shading(cell, "002E6D")
    cell_last = row.cells[-1]
    add_form_field_to_cell(cell_last, field_name="pl_final_qty",
                           size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(cell_last, "F2F2F2")
    set_cell_border(cell_last, bottom={"sz": "4", "color": "002E6D"})

    # Column widths
    widths = [Inches(1.3), Inches(2.0), Inches(1.0), Inches(1.0), Inches(1.2)]
    for row in table.rows:
        for j, w in enumerate(widths):
            if j < len(row.cells):
                row.cells[j].width = w


def build_product_pages(doc, products_folder, lang="en"):
    """Photo pages for each product subfolder."""
    products_path = Path(products_folder)
    if not products_path.exists():
        return

    subfolders = sorted(
        [d for d in products_path.iterdir() if d.is_dir()],
        key=lambda d: natural_sort_key(d.name)
    )

    first_item = True
    for folder in subfolders:
        images = get_sorted_images(folder)
        if not images:
            continue

        item_name = folder.name.upper()
        doc.add_page_break()

        if first_item:
            add_section_heading(doc, t("sec_product_details", lang))
            first_item = False

        # Compact item header (matches Globex format):
        # Row 1: Item no. [NAME]  |  RESULT  |  COMMENTS
        # Row 2: Label:           |  ___     |  ___
        info_table = doc.add_table(rows=2, cols=3)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Row 1: header
        cell0 = info_table.rows[0].cells[0]
        p0 = cell0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(2)
        p0.paragraph_format.space_after = Pt(2)
        run = p0.add_run("Item no.  ")
        run.font.name = FONT_NAME
        run.font.size = Pt(8)
        run.font.color.rgb = NAVY
        run2 = p0.add_run(item_name)
        run2.font.name = FONT_NAME
        run2.font.size = Pt(10)
        run2.font.bold = True
        run2.font.color.rgb = NAVY
        set_cell_shading(cell0, "D6E4F0")

        fmt(info_table.rows[0].cells[1], t("result", lang),
            size=Pt(8), bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
        fmt(info_table.rows[0].cells[2], t("chk_header_comments", lang),
            size=Pt(8), bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Row 2: label check — fillable
        fmt(info_table.rows[1].cells[0], t("label", lang),
            size=Pt(9), bold=False, color=CHARCOAL)
        set_cell_shading(info_table.rows[1].cells[1], "F2F2F2")
        add_form_field_to_cell(info_table.rows[1].cells[1],
                               field_name=f"item_{item_name}_result",
                               size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(info_table.rows[1].cells[2], "F2F2F2")
        add_form_field_to_cell(info_table.rows[1].cells[2],
                               field_name=f"item_{item_name}_comments",
                               size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER)

        # Column widths
        info_table.rows[0].cells[0].width = Inches(4.0)
        info_table.rows[0].cells[1].width = Inches(1.0)
        info_table.rows[0].cells[2].width = Inches(1.8)
        # Thin borders
        for row in info_table.rows:
            for cell in row.cells:
                set_cell_border(cell, bottom={"sz": "2", "color": "D9D9D9"})

        # Photo grid
        _add_photo_grid(doc, images, lang=lang)


def _add_photo_grid(doc, images, caption_labels=None, lang="en"):
    """Insert images in a 2-column grid, 6 per page (3 rows of 2)."""
    PHOTOS_PER_PAGE = 6
    IMG_WIDTH = Inches(3.3)
    IMG_MAX_H = 1.9  # inches — fits 3 rows + captions + margins on letter

    for i in range(0, len(images), 2):
        batch = images[i:i + 2]

        # Photo row
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        make_table_borderless(table)

        for j, img_path in enumerate(batch):
            cell = table.rows[0].cells[j]
            set_cell_margins(cell, top=20, bottom=10, left=30, right=30)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            try:
                buf, (w, h) = resize_image_for_doc(
                    img_path, max_width_inches=3.3, max_height_inches=IMG_MAX_H)
                run = cell.paragraphs[0].add_run()
                run.add_picture(buf, width=IMG_WIDTH)
            except Exception as e:
                run = cell.paragraphs[0].add_run(f"[Image error: {e}]")
                run.font.size = Pt(7)

        # Caption row
        cap_table = doc.add_table(rows=1, cols=2)
        cap_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        make_table_borderless(cap_table)
        for j in range(min(2, len(batch))):
            cell = cap_table.rows[0].cells[j]
            set_cell_shading(cell, "002E6D")
            set_cell_margins(cell, top=10, bottom=10, left=30, right=30)
            label = t("product_caption", lang)
            if caption_labels and i + j < len(caption_labels):
                label = caption_labels[i + j]
            fmt(cell, label, size=Pt(7), color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        if len(batch) == 1:
            set_cell_shading(cap_table.rows[0].cells[1], "FFFFFF")

        # Small spacer between rows (not after last pair)
        if (i + 2) < len(images):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)


def build_general_product_pictures(doc, products_folder, lang="en"):
    """General product pictures — representative photos from each size."""
    products_path = Path(products_folder)
    if not products_path.exists():
        return

    subfolders = sorted(
        [d for d in products_path.iterdir() if d.is_dir()],
        key=lambda d: natural_sort_key(d.name)
    )

    images = []
    captions = []
    for folder in subfolders:
        folder_images = get_sorted_images(folder)
        if not folder_images:
            continue
        name = folder.name.upper()
        selected = folder_images[-2:] if len(folder_images) >= 2 else folder_images
        for img in selected:
            images.append(img)
            captions.append(name)

    if not images:
        return

    doc.add_page_break()
    add_subsection_heading(doc, t("general_pictures", lang))
    doc.add_paragraph()
    _add_photo_grid(doc, images, caption_labels=captions, lang=lang)


def build_container_page(doc, container_folder, lang="en"):
    """Container details table + container photos."""
    doc.add_page_break()
    add_section_heading(doc, t("sec_container", lang))

    # Container info table
    headers = [
        t("ct_type", lang),
        t("ct_number", lang),
        t("ct_seal", lang),
        t("ct_loading_qty", lang),
    ]
    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        set_cell_shading(cell, "002E6D")
        fmt(cell, header, size=Pt(9), bold=True, color=WHITE,
            align=WD_ALIGN_PARAGRAPH.CENTER)

    ct_field_names = ["ct_type", "ct_number", "ct_seal", "ct_loading_qty"]
    for j in range(4):
        cell = table.rows[1].cells[j]
        set_cell_shading(cell, "F2F2F2")
        add_form_field_to_cell(cell, field_name=ct_field_names[j],
                               size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_border(cell, bottom={"sz": "2", "color": "D9D9D9"})

    doc.add_paragraph()

    # Container photos
    container_path = Path(container_folder)
    if container_path.exists():
        images = get_sorted_images(container_path)
        if images:
            captions = [
                t("cap_container_num", lang),
                t("cap_seal", lang),
                t("cap_empty", lang),
                t("cap_light_test", lang),
            ]
            while len(captions) < len(images):
                captions.append(t("cap_container", lang))
            _add_photo_grid(doc, images, caption_labels=captions, lang=lang)


def build_container_damages(doc, lang="en"):
    """Container damages section."""
    doc.add_page_break()
    add_subsection_heading(doc, t("container_damages", lang))
    add_placeholder(doc, t("placeholder_damages", lang))

    add_subsection_heading(doc, t("remarks", lang))
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    add_form_field(p, field_name="remarks_container", size=Pt(9), color=PLACEHOLDER_GRAY)


def build_loading_process(doc, lang="en"):
    """Loading process with 25/50/75/100% placeholders."""
    doc.add_page_break()
    add_section_heading(doc, t("sec_loading", lang))

    captions = ["25%", "50%", "75%", "100%"]
    for row_idx in range(2):
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        make_table_borderless(table)
        for col_idx in range(2):
            cell = table.rows[0].cells[col_idx]
            set_cell_margins(cell, top=40, bottom=20, left=40, right=40)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Gray placeholder box
            set_cell_shading(cell, "F2F2F2")
            run = cell.paragraphs[0].add_run(f"\n\n{t('placeholder_photo', lang)}\n\n")
            run.font.name = FONT_NAME
            run.font.size = Pt(9)
            run.font.color.rgb = PLACEHOLDER_GRAY

        cap_table = doc.add_table(rows=1, cols=2)
        cap_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        make_table_borderless(cap_table)
        for col_idx in range(2):
            cell = cap_table.rows[0].cells[col_idx]
            set_cell_shading(cell, "002E6D")
            set_cell_margins(cell, top=20, bottom=20, left=40, right=40)
            fmt(cell, captions[row_idx * 2 + col_idx], size=Pt(7), color=WHITE,
                align=WD_ALIGN_PARAGRAPH.CENTER)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)


def build_acknowledgment_page(doc, docs_dir=None, other_dir=None, lang="en"):
    """CCI Acknowledgment — Trifecta-branded field form, single page."""
    doc.add_page_break()

    # --- Title (bilingual) ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(t("ack_title", lang) + "  " + t("ack_title_zh", lang))
    run.font.name = FONT_NAME
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = NAVY

    # --- Data fields (2-column layout, compact) ---
    field_pairs = [
        (t("ack_date", lang), t("ack_loading_time", lang)),
        (t("ack_container_no", lang), t("ack_reference_no", lang)),
        (t("ack_product", lang), t("ack_seal_no", lang)),
        (t("ack_items", lang), t("ack_sample_standard", lang)),
        (t("ack_package_qty", lang), ""),
    ]

    table = doc.add_table(rows=len(field_pairs), cols=4)
    make_table_borderless(table)
    for i, (left_label, right_label) in enumerate(field_pairs):
        row = table.rows[i]
        c0 = row.cells[0]
        fmt(c0, left_label, size=Pt(7), bold=True, color=CHARCOAL)
        c0.paragraphs[0].paragraph_format.space_before = Pt(0)
        c0.paragraphs[0].paragraph_format.space_after = Pt(0)
        c0.width = Inches(1.6)

        c1 = row.cells[1]
        set_cell_border(c1, bottom={"sz": "2", "color": "D9D9D9"})
        c1.width = Inches(1.6)
        add_form_field_to_cell(c1, size=Pt(8), color=CHARCOAL)
        c1.paragraphs[0].paragraph_format.space_before = Pt(0)
        c1.paragraphs[0].paragraph_format.space_after = Pt(0)

        c2 = row.cells[2]
        c2.width = Inches(1.6)
        if right_label:
            fmt(c2, right_label, size=Pt(7), bold=True, color=CHARCOAL)
            c2.paragraphs[0].paragraph_format.space_before = Pt(0)
            c2.paragraphs[0].paragraph_format.space_after = Pt(0)

        c3 = row.cells[3]
        c3.width = Inches(1.6)
        if right_label:
            set_cell_border(c3, bottom={"sz": "2", "color": "D9D9D9"})
            add_form_field_to_cell(c3, size=Pt(8), color=CHARCOAL)
            c3.paragraphs[0].paragraph_format.space_before = Pt(0)
            c3.paragraphs[0].paragraph_format.space_after = Pt(0)

    # --- Checklist (compact) ---
    checklist_items = [
        t("ack_photos_allowed", lang),
        t("ack_labels", lang),
        t("ack_fumigation", lang),
        t("ack_qty_check", lang),
        t("ack_container_quality", lang),
        t("ack_box_quality", lang),
        t("ack_all_items_photo", lang),
        t("ack_space_enough", lang),
        t("ack_same_qty_pl", lang),
    ]

    chk_table = doc.add_table(rows=1 + len(checklist_items), cols=2)
    chk_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    set_cell_shading(chk_table.rows[0].cells[0], "002E6D")
    fmt(chk_table.rows[0].cells[0], t("chk_header_item", lang),
        size=Pt(7), bold=True, color=WHITE)
    chk_table.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Pt(1)
    chk_table.rows[0].cells[0].paragraphs[0].paragraph_format.space_after = Pt(1)
    set_cell_shading(chk_table.rows[0].cells[1], "002E6D")
    fmt(chk_table.rows[0].cells[1], t("chk_header_result", lang),
        size=Pt(7), bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    chk_table.rows[0].cells[1].paragraphs[0].paragraph_format.space_before = Pt(1)
    chk_table.rows[0].cells[1].paragraphs[0].paragraph_format.space_after = Pt(1)

    for i, item in enumerate(checklist_items):
        row = chk_table.rows[i + 1]
        fmt(row.cells[0], item, size=Pt(7), color=CHARCOAL)
        row.cells[0].paragraphs[0].paragraph_format.space_before = Pt(0)
        row.cells[0].paragraphs[0].paragraph_format.space_after = Pt(0)
        set_cell_shading(row.cells[1], "F2F2F2")
        add_form_field_to_cell(row.cells[1], size=Pt(7),
                               align=WD_ALIGN_PARAGRAPH.CENTER)
        row.cells[1].paragraphs[0].paragraph_format.space_before = Pt(0)
        row.cells[1].paragraphs[0].paragraph_format.space_after = Pt(0)
        for j in range(2):
            set_cell_border(row.cells[j], bottom={"sz": "1", "color": "D9D9D9"})

    chk_table.rows[0].cells[0].width = Inches(4.8)
    chk_table.rows[0].cells[1].width = Inches(1.2)

    # --- Inspection result (inline) ---
    result_table = doc.add_table(rows=1, cols=3)
    result_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_table_borderless(result_table)

    # Label
    cell_label = result_table.rows[0].cells[0]
    cell_label.width = Inches(2.5)
    fmt(cell_label, t("ack_inspection_result", lang),
        size=Pt(10), bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.RIGHT)
    cell_label.paragraphs[0].paragraph_format.space_before = Pt(3)
    cell_label.paragraphs[0].paragraph_format.space_after = Pt(3)

    for j, (label, color_hex) in enumerate([
        (t("ack_pass", lang), "28A745"),
        (t("ack_fail", lang), "DC3545"),
    ]):
        cell = result_table.rows[0].cells[j + 1]
        cell.width = Inches(1.5)
        set_cell_shading(cell, "F8F8F8")
        set_cell_border(cell,
                        top={"sz": "4", "color": color_hex},
                        bottom={"sz": "4", "color": color_hex},
                        left={"sz": "4", "color": color_hex},
                        right={"sz": "4", "color": color_hex})
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"\u2610  {label}")
        run.font.name = FONT_NAME
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(
            int(color_hex[0:2], 16),
            int(color_hex[2:4], 16),
            int(color_hex[4:6], 16),
        )

    # --- Disclaimer (compact) ---
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(t("ack_disclaimer", lang))
    run.font.name = FONT_NAME
    run.font.size = Pt(5.5)
    run.font.italic = True
    run.font.color.rgb = MID_GRAY

    # --- Comments (single line) ---
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(t("ack_comments", lang) + "  ")
    run.font.name = FONT_NAME
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = NAVY
    add_form_field(p, field_name="ack_comments", size=Pt(8), color=CHARCOAL)

    # --- Signature lines (compact) ---
    sig_table = doc.add_table(rows=2, cols=2)
    make_table_borderless(sig_table)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j in range(2):
        cell = sig_table.rows[0].cells[j]
        cell.width = Inches(3.0)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(0)
        set_cell_border(cell, bottom={"sz": "4", "color": "002E6D"})

    labels = [t("ack_inspector", lang), t("ack_supplier", lang)]
    for j, label in enumerate(labels):
        fmt(sig_table.rows[1].cells[j], label,
            size=Pt(8), bold=True, color=NAVY,
            align=WD_ALIGN_PARAGRAPH.CENTER)
        sig_table.rows[1].cells[j].paragraphs[0].paragraph_format.space_before = Pt(2)
        sig_table.rows[1].cells[j].paragraphs[0].paragraph_format.space_after = Pt(0)


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------
def generate_report(folder_path, output_path=None, lang="en"):
    """Generate a CCI inspection report from an inspection folder.

    Args:
        folder_path: Path to the inspection folder (must contain 2.Photos/)
        output_path: Where to save the .docx
        lang: Language code — "en" (English) or "es" (Spanish)
    """
    folder = Path(folder_path)
    if not folder.exists():
        raise FileNotFoundError(f"Folder not found: {folder}")

    photos_dir = folder / "2.Photos"
    products_dir = photos_dir / "1.Products"
    container_dir = photos_dir / "2.Container"

    if not photos_dir.exists():
        raise FileNotFoundError(f"2.Photos/ folder not found in {folder}")

    if output_path is None:
        output_path = folder / "report_output.docx"

    print(f"Generating report from: {folder}")
    print(f"Output: {output_path}")

    doc = create_document()
    add_header_footer(doc)

    print(f"  Language: {lang}")

    print("  [1/7] Summary page...")
    build_page1_summary(doc, lang)

    print("  [2/7] Packing list...")
    build_page2_packing_list(doc, lang)

    print("  [3/7] Product photos...")
    build_product_pages(doc, products_dir, lang)

    print("  [4/7] General product pictures...")
    build_general_product_pictures(doc, products_dir, lang)

    print("  [5/7] Container details...")
    build_container_page(doc, container_dir, lang)
    build_container_damages(doc, lang)

    print("  [6/7] Loading process...")
    build_loading_process(doc, lang)

    print("  [7/7] Acknowledgment...")
    other_dir = photos_dir / "Other"
    docs_dir = folder / "3.Documents"
    build_acknowledgment_page(doc, docs_dir, other_dir, lang)

    doc.save(str(output_path))
    print(f"\nDone: {output_path}")
    return str(output_path)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate Trifecta CCI Inspection Report")
    parser.add_argument("folder", help="Path to the inspection folder")
    parser.add_argument("-o", "--output", help="Output .docx path", default=None)
    parser.add_argument("-l", "--lang", choices=["en", "es"], default="en",
                        help="Report language: en (English) or es (Spanish)")
    args = parser.parse_args()
    generate_report(args.folder, args.output, lang=args.lang)
